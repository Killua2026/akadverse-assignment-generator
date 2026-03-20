"""
AkadVerse — Assignment Generator
Tier 5 | Microservice Port: 8008
========================================================================
v1.0 — Initial build.

What this service does:
  A faculty-facing tool. A lecturer provides their course title, topic,
  academic level, learning outcomes, and assignment details. Gemini
  generates three structured academic documents in a single API call:

    1. Assignment Brief  — the document given to students describing
                           the task, requirements, format, and deadline.
    2. Marking Rubric    — a criteria grid with descriptors for each
                           grade band (Distinction, Merit, Pass, Fail).
    3. Marking Scheme    — model answers and point allocations used by
                           the lecturer when grading submissions.

  Output pipeline (local dev simulation):
    - All three documents are saved as .docx files in OUTPUT_DIR
      (simulates Google Cloud Storage).
    - Metadata is logged to SQLite (simulates PostgreSQL).
    - Optionally synced to Google Docs via the Docs API (real, using
      credentials from Tier 4 Google Workspace Integration).
    - A Kafka mock event 'assignment.created' is published to the bus.

Architecture note:
  This is a pure LLM generation tool — no RAG, no vector stores, no
  scraping. The entire complexity lives in the Pydantic output schema
  and the prompt. Gemini generates all three documents in one structured
  output call using with_structured_output(method="json_mode").
"""

import os
import json
import sqlite3
import threading
from datetime import datetime
from contextlib import asynccontextmanager, contextmanager
from typing import AsyncIterator, Dict, List, Optional, Union
from uuid import uuid4

from fastapi import FastAPI, HTTPException, Form
from pydantic import BaseModel, Field

# Unified Google GenAI SDK (v1.67.0+)
from google import genai

# LangChain for structured output generation
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate

# python-docx for local .docx file generation (simulates GCS storage)
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =========================================================
# CONSTANTS
# =========================================================

# Local directory where generated .docx files are saved.
# Simulates Google Cloud Storage in the development environment.
OUTPUT_DIR = "generated_assignments"

# SQLite database path for assignment metadata.
# Simulates PostgreSQL in the development environment.
DB_PATH = "akadverse_assignments.db"

# Maximum characters for learning outcomes input to keep prompt size
# manageable and avoid exceeding Gemini's context for structured output.
MAX_LEARNING_OUTCOMES_LENGTH = 2000

# Threading lock for SQLite writes — prevents corruption if two lecturers
# submit generation requests concurrently.
db_lock = threading.Lock()


# =========================================================
# PYDANTIC OUTPUT SCHEMAS
# =========================================================

class RubricRow(BaseModel):
    """
    A single row in the marking rubric grid.
    Each row represents one assessment criterion (e.g. 'Critical Analysis')
    with descriptors for each of the four grade bands.
    """
    criterion: str = Field(description="The assessment criterion, e.g. 'Critical Analysis and Argument'.")
    distinction: str = Field(description="Descriptor for Distinction (70-100%): what excellent work looks like for this criterion.")
    merit: str = Field(description="Descriptor for Merit (60-69%): what good work looks like for this criterion.")
    pass_band: str = Field(description="Descriptor for Pass (40-59%): what satisfactory work looks like for this criterion.")
    fail: str = Field(description="Descriptor for Fail (0-39%): what inadequate work looks like for this criterion.")
    marks_available: int = Field(description="Total marks available for this criterion.")


class MarkingSchemeItem(BaseModel):
    """
    A single item in the marking scheme — one question or task component
    with its model answer and how marks are allocated.
    """
    task_number: str = Field(description="Task or question identifier, e.g. 'Task 1', 'Part (a)'.")
    task_description: str = Field(description="Brief restatement of what this task asks students to do.")
    model_answer: str = Field(description="A thorough model answer or list of key points expected from students.")
    marks: int = Field(description="Marks allocated to this task or question.")
    marking_guidance: str = Field(description="Specific guidance for the marker: what to award full/partial marks for, common errors to watch out for.")


class AssignmentBrief(BaseModel):
    """
    The student-facing assignment document. Describes what students must
    do, how to do it, the format requirements, and submission details.
    """
    title: str = Field(description="Full assignment title, e.g. 'CSC 301 Individual Assignment: Algorithm Complexity Analysis'.")
    course_code_and_title: str = Field(description="e.g. 'CSC 301 — Data Structures and Algorithms'.")
    academic_level: str = Field(description="e.g. '300 Level, Second Semester'.")
    total_marks: int = Field(description="Total marks this assignment is worth.")
    weighting: str = Field(description="Percentage of final grade, e.g. '20% of final grade'.")
    submission_deadline: str = Field(description="Submission deadline as provided or a suggested reasonable timeframe.")
    background: str = Field(description="1-2 paragraphs providing academic context and rationale for the assignment.")
    learning_outcomes_assessed: List[str] = Field(description="Bullet list of the specific learning outcomes this assignment tests.")
    tasks: List[str] = Field(description="Numbered list of the specific tasks students must complete, each described clearly and unambiguously.")
    format_requirements: str = Field(description="Word count, file format, font, referencing style, and any other submission format rules.")
    submission_instructions: str = Field(description="How and where to submit — e.g. 'Upload to the AkadVerse platform under Assignments tab'.")
    academic_integrity_note: str = Field(description="Standard academic integrity and plagiarism warning appropriate for the institution.")


class MarkingRubric(BaseModel):
    """
    The assessment rubric used to grade student submissions consistently.
    Contains one row per criterion, each with descriptors for four grade bands.
    """
    title: str = Field(description="Rubric title, e.g. 'Marking Rubric: Algorithm Complexity Analysis Assignment'.")
    grade_boundaries: str = Field(description="Grade band boundaries used in this rubric, e.g. 'Distinction 70-100% | Merit 60-69% | Pass 40-59% | Fail 0-39%'.")
    criteria: List[RubricRow] = Field(description="List of assessment criteria rows. Include 4-6 criteria covering the key aspects of the assignment.")
    total_marks: int = Field(description="Sum of all marks_available across all criteria — must match the brief's total_marks.")
    general_marking_guidance: str = Field(description="Overall guidance for markers: consistency notes, handling borderline cases, etc.")


class MarkingScheme(BaseModel):
    """
    The lecturer's marking scheme — model answers and mark allocations
    for every task in the assignment brief.
    """
    title: str = Field(description="Scheme title, e.g. 'Marking Scheme: Algorithm Complexity Analysis Assignment'.")
    items: List[MarkingSchemeItem] = Field(description="One item per task from the brief, with model answer and marks.")
    total_marks: int = Field(description="Total marks — must match the brief and rubric.")
    marker_notes: str = Field(description="Additional notes for the marker: borderline decisions, acceptable alternative answers, common pitfalls.")


class AssignmentPackage(BaseModel):
    """
    The root schema returned by Gemini in a single structured output call.
    Contains all three documents as nested objects.
    """
    brief: AssignmentBrief = Field(description="The student-facing assignment brief.")
    rubric: MarkingRubric = Field(description="The marking rubric with grade band descriptors.")
    scheme: MarkingScheme = Field(description="The marking scheme with model answers and mark allocations.")


class GenerationSuccessResponse(BaseModel):
    """Schema for the API success response returned to the caller."""
    status: str
    assignment_id: str
    message: str
    course_title: str
    topic: str
    files_saved: List[str]
    docs_synced: bool
    total_marks: int


# =========================================================
# DATABASE HELPERS
# =========================================================

def init_db() -> None:
    """
    Creates the SQLite assignments table on startup if it does not exist.
    The table stores metadata for every generated assignment package.
    This simulates the PostgreSQL metadata store specified in the architecture.
    """
    try:
        with get_db_connection() as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS assignments (
                    id              TEXT    PRIMARY KEY,
                    course_title    TEXT    NOT NULL,
                    topic           TEXT    NOT NULL,
                    academic_level  TEXT    NOT NULL,
                    total_marks     INTEGER NOT NULL,
                    brief_path      TEXT    NOT NULL,
                    rubric_path     TEXT    NOT NULL,
                    scheme_path     TEXT    NOT NULL,
                    docs_synced     INTEGER NOT NULL DEFAULT 0,
                    created_at      TEXT    NOT NULL
                )
            """)
            conn.commit()
        print("[DB] Assignments database initialised successfully.")
    except sqlite3.Error as e:
        print(f"[DB ERROR] Initialisation failed: {e}")
        raise


@contextmanager
def get_db_connection():
    """
    Context manager for SQLite connections. Guarantees the connection is
    closed even if an exception fires mid-operation, preventing leaks.
    Row factory is set to sqlite3.Row so columns are accessible by name.
    """
    conn = None
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        yield conn
    except sqlite3.Error as e:
        if conn:
            conn.rollback()
        print(f"[DB ERROR] {e}")
        raise
    finally:
        if conn:
            conn.close()


def log_assignment_to_db(
    assignment_id: str,
    course_title: str,
    topic: str,
    academic_level: str,
    total_marks: int,
    brief_path: str,
    rubric_path: str,
    scheme_path: str,
    docs_synced: bool
) -> None:
    """
    Persists assignment generation metadata to SQLite.
    Called after all three documents have been saved successfully.
    Uses db_lock to prevent concurrent write corruption.
    """
    with db_lock:
        with get_db_connection() as conn:
            conn.execute(
                """
                INSERT INTO assignments
                  (id, course_title, topic, academic_level, total_marks,
                   brief_path, rubric_path, scheme_path, docs_synced, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    assignment_id,
                    course_title,
                    topic,
                    academic_level,
                    total_marks,
                    brief_path,
                    rubric_path,
                    scheme_path,
                    int(docs_synced),
                    datetime.utcnow().isoformat()
                )
            )
            conn.commit()
    print(f"[DB] Assignment '{assignment_id}' logged to database.")


# =========================================================
# MODEL DISCOVERY
# =========================================================

def get_valid_model_name(api_key_str: str) -> str:
    """
    Dynamically discovers the best available Gemini generative model
    by calling client.models.list() and matching against a priority list.

    Consistent with the same pattern used across all Tier 5 microservices.
    Falls back to 'gemini-1.5-flash' if discovery fails entirely.
    """
    try:
        client = genai.Client(api_key=api_key_str)

        all_models = [
            m.name.replace("models/", "")
            for m in client.models.list()
            if m.name
        ]

        priority_order = [
            "gemini-2.5-flash",
            "gemini-2.0-flash",
            "gemini-1.5-flash",
            "gemini-pro",
        ]
        for preferred in priority_order:
            if preferred in all_models:
                print(f"[Model] Using: {preferred}")
                return preferred

        if all_models:
            print(f"[Model] Fallback to first available: {all_models[0]}")
            return all_models[0]

    except Exception as e:
        print(f"[Model WARNING] Discovery failed ({e}). Defaulting to 'gemini-1.5-flash'.")

    return "gemini-1.5-flash"


# =========================================================
# DOCUMENT GENERATION HELPERS
# =========================================================

def save_brief_as_docx(brief: AssignmentBrief, path: str) -> None:
    """
    Saves the assignment brief as a formatted .docx file.
    Uses python-docx to produce a document that looks like a real
    university assignment brief -- not a plain text dump.

    Simulates saving to Google Cloud Storage (GCS) in production.
    """
    doc = DocxDocument()

    # ---- Title block ----
    title_para = doc.add_heading(brief.title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Course: {brief.course_code_and_title}")
    doc.add_paragraph(f"Level: {brief.academic_level}")
    doc.add_paragraph(f"Total Marks: {brief.total_marks}  |  Weighting: {brief.weighting}")
    doc.add_paragraph(f"Submission Deadline: {brief.submission_deadline}")
    doc.add_paragraph("")

    # ---- Background ----
    doc.add_heading("Background", level=2)
    doc.add_paragraph(brief.background)

    # ---- Learning outcomes ----
    doc.add_heading("Learning Outcomes Assessed", level=2)
    for outcome in brief.learning_outcomes_assessed:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(outcome)

    # ---- Tasks ----
    doc.add_heading("Assignment Tasks", level=2)
    for i, task in enumerate(brief.tasks, start=1):
        doc.add_paragraph(f"{i}. {task}")

    # ---- Format requirements ----
    doc.add_heading("Format Requirements", level=2)
    doc.add_paragraph(brief.format_requirements)

    # ---- Submission instructions ----
    doc.add_heading("Submission Instructions", level=2)
    doc.add_paragraph(brief.submission_instructions)

    # ---- Academic integrity ----
    doc.add_heading("Academic Integrity", level=2)
    doc.add_paragraph(brief.academic_integrity_note)

    doc.save(path)
    print(f"[DocX] Brief saved: {path}")


def save_rubric_as_docx(rubric: MarkingRubric, path: str) -> None:
    """
    Saves the marking rubric as a formatted .docx table.
    Creates a proper rubric grid: criteria in rows, grade bands in columns.
    This format matches what lecturers at Nigerian universities typically use.
    """
    doc = DocxDocument()

    doc.add_heading(rubric.title, level=1)
    doc.add_paragraph(f"Grade Boundaries: {rubric.grade_boundaries}")
    doc.add_paragraph("")

    # ---- Rubric table ----
    # Columns: Criterion | Distinction | Merit | Pass | Fail | Marks
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    # Header row
    headers = ["Criterion", "Distinction (70-100%)", "Merit (60-69%)", "Pass (40-59%)", "Fail (0-39%)", "Marks"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Bold the header text
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows — one per criterion
    for row_data in rubric.criteria:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data.criterion
        row_cells[1].text = row_data.distinction
        row_cells[2].text = row_data.merit
        row_cells[3].text = row_data.pass_band
        row_cells[4].text = row_data.fail
        row_cells[5].text = str(row_data.marks_available)

    doc.add_paragraph("")
    doc.add_paragraph(f"Total Marks: {rubric.total_marks}")
    doc.add_heading("General Marking Guidance", level=2)
    doc.add_paragraph(rubric.general_marking_guidance)

    doc.save(path)
    print(f"[DocX] Rubric saved: {path}")


def save_scheme_as_docx(scheme: MarkingScheme, path: str) -> None:
    """
    Saves the marking scheme as a formatted .docx file.
    Each task gets its own section with model answer, marks, and
    specific guidance for the marker.
    """
    doc = DocxDocument()

    doc.add_heading(scheme.title, level=1)
    doc.add_paragraph(f"Total Marks: {scheme.total_marks}")
    doc.add_paragraph("")

    # ---- One section per task ----
    for item in scheme.items:
        doc.add_heading(f"{item.task_number}: {item.task_description}", level=2)

        doc.add_paragraph("Model Answer:", style="Intense Quote")
        doc.add_paragraph(item.model_answer)

        marks_para = doc.add_paragraph()
        marks_run = marks_para.add_run(f"Marks: {item.marks}")
        marks_run.bold = True

        doc.add_paragraph("Marking Guidance:")
        doc.add_paragraph(item.marking_guidance)
        doc.add_paragraph("")

    # ---- Marker notes ----
    doc.add_heading("Marker Notes", level=2)
    doc.add_paragraph(scheme.marker_notes)

    doc.save(path)
    print(f"[DocX] Scheme saved: {path}")


# =========================================================
# GOOGLE DOCS SYNC (optional — uses Tier 4 credentials)
# =========================================================

def sync_to_google_docs(
    brief: AssignmentBrief,
    rubric: MarkingRubric,
    scheme: MarkingScheme,
    assignment_id: str,
    google_api_key: str,
    brief_path: str,
    rubric_path: str,
    scheme_path: str
) -> bool:
    """
    Uploads the three generated .docx files to Google Drive, placing them
    inside the /AkadVerse/2026/Assignments/ folder structure.

    This mirrors the pattern used by the Tier 4 Google Workspace Integration
    (drive_handler.py) and uploads the ACTUAL .docx files rather than
    recreating them as truncated plain text -- so the Drive copies are
    identical in quality to the locally saved files.

    The upload uses Drive API's MIME type conversion:
    uploading a .docx with mimeType='application/vnd.google-apps.document'
    tells Google to convert it into a native Google Doc on arrival,
    making it fully editable in the browser.

    Token expiry note: if this function returns False with an
    'invalid_grant' error in the terminal, your token.json has expired.
    Fix: navigate to your akadverse-workspace-service folder, run
    `uvicorn main:app --host 127.0.0.1 --port 8002 --reload`, visit
    http://127.0.0.1:8002/login in your browser, complete the Google
    sign-in, then copy the fresh token.json back into this folder.

    Returns True if all three uploads succeeded, False otherwise.
    """
    token_path = "token.json"

    def log_docs_sync(level: str, message: str, **context: object) -> None:
        """Structured logger for docs sync operations with assignment-level context."""
        payload = {"assignment_id": assignment_id, **context}
        try:
            context_text = json.dumps(payload, default=str)
        except Exception:
            context_text = str(payload)
        print(f"[Docs Sync {level}] {message} | context={context_text}")

    # Guard: check for credentials before attempting any API call
    if not os.path.exists(token_path):
        log_docs_sync(
            "WARNING",
            "No token.json found. Skipping Google Docs sync. "
            "Copy token.json from akadverse-workspace-service to enable.",
            token_path=token_path
        )
        return False

    try:
        import io
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request as GoogleRequest
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaIoBaseUpload

        # Load credentials from the Tier 4 token file
        scopes = [
            "https://www.googleapis.com/auth/drive.file",
            "https://www.googleapis.com/auth/documents"
        ]
        creds = Credentials.from_authorized_user_file(token_path, scopes)

        # Attempt to refresh an expired-but-refreshable token automatically.
        # This handles the common case where the access token (1 hour TTL)
        # has expired but the refresh token is still valid.
        # 'invalid_grant' means the refresh token itself has expired --
        # that requires a manual re-authentication (see token expiry note above).
        if not creds.valid:
            if creds.expired and creds.refresh_token:
                log_docs_sync("INFO", "Access token expired. Refreshing automatically...")
                creds.refresh(GoogleRequest())
                # Persist the refreshed token so future calls don't need to refresh
                with open(token_path, "w") as f:
                    f.write(creds.to_json())
                log_docs_sync("INFO", "Token refreshed and saved.", token_path=token_path)
            else:
                log_docs_sync(
                    "WARNING",
                    "Token is invalid and cannot be refreshed. "
                    "Re-authenticate via akadverse-workspace-service /login endpoint, "
                    "then copy the fresh token.json here.",
                    token_path=token_path
                )
                return False

        drive_service = build("drive", "v3", credentials=creds)

        # ---- Step 1: Ensure /AkadVerse/2026/Assignments/ folder structure exists ----
        # Mirrors the get_or_create_folder pattern from drive_handler.py
        def get_or_create_folder(name: str, parent_id: Optional[str] = None) -> str:
            """
            Checks if a Drive folder exists by name (optionally under a parent).
            Creates it if missing. Returns the folder ID.
            Raises RuntimeError if the API call fails.
            """
            query = (
                f"name='{name}' and "
                f"mimeType='application/vnd.google-apps.folder' and "
                f"trashed=false"
            )
            if parent_id:
                query += f" and '{parent_id}' in parents"

            results = drive_service.files().list(
                q=query, spaces="drive", fields="files(id, name)"
            ).execute()

            if not isinstance(results, dict):
                raise RuntimeError("Drive API list response was not a dictionary.")

            files = results.get("files", [])
            if not isinstance(files, list):
                raise RuntimeError("Drive API list response contained a non-list 'files' field.")
            if files:
                first_file = files[0]
                if not isinstance(first_file, dict):
                    raise RuntimeError("Drive API list response contained an invalid file record.")

                existing_id = first_file.get("id")
                if not existing_id:
                    raise RuntimeError(
                        f"Drive API returned folder '{name}' without an 'id' field."
                    )
                log_docs_sync(
                    "INFO",
                    "Drive folder found.",
                    folder_name=name,
                    parent_id=parent_id,
                    folder_id=existing_id
                )
                return existing_id

            # Folder does not exist -- create it
            # 'parents' is optional and becomes list[str] when provided,
            # so the value type must support both str and list[str].
            metadata: Dict[str, Union[str, List[str]]] = {
                "name": name,
                "mimeType": "application/vnd.google-apps.folder",
            }
            if parent_id:
                metadata["parents"] = [parent_id]

            folder = drive_service.files().create(
                body=metadata, fields="id"
            ).execute()
            if not isinstance(folder, dict):
                raise RuntimeError("Drive API folder creation response was not a dictionary.")

            new_id = folder.get("id")
            if not new_id:
                raise RuntimeError(
                    f"Drive API folder creation for '{name}' succeeded without returning an ID."
                )
            log_docs_sync(
                "INFO",
                "Drive folder created.",
                folder_name=name,
                parent_id=parent_id,
                folder_id=new_id
            )
            return new_id

        # Build /AkadVerse/2026/Assignments/ -- consistent with Tier 4's Notes structure
        root_id        = get_or_create_folder("AkadVerse")
        year_id        = get_or_create_folder("2026", parent_id=root_id)
        target_id      = get_or_create_folder("Assignments", parent_id=year_id)

        # ---- Step 2: Upload each .docx file as a native Google Doc ----
        # Setting mimeType='application/vnd.google-apps.document' on the
        # file metadata triggers Drive's automatic .docx → Google Doc conversion,
        # so the file lands as a fully editable document, not a binary attachment.
        def upload_docx_as_google_doc(docx_path: str, doc_title: str, folder_id: str) -> str:
            """
            Uploads a local .docx file to Drive and converts it to a Google Doc.
            Returns the webViewLink (browser URL) of the created document.
            Raises an exception if the upload fails.
            """
            if not os.path.exists(docx_path):
                raise FileNotFoundError(f"Source file not found: {docx_path}")
            if not folder_id.strip():
                raise ValueError("Target Drive folder ID is empty.")

            file_metadata = {
                "name": doc_title,
                # This MIME type tells Drive to convert the upload to a Google Doc
                "mimeType": "application/vnd.google-apps.document",
                "parents": [folder_id]
            }

            # Read the .docx bytes and wrap in a file-like object for the API
            with open(docx_path, "rb") as f:
                docx_bytes = f.read()
            if not docx_bytes:
                raise RuntimeError(f"Source .docx file is empty: {docx_path}")

            media = MediaIoBaseUpload(
                io.BytesIO(docx_bytes),
                # Source MIME type is the .docx format
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                resumable=True
            )

            uploaded = drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields="id, webViewLink"
            ).execute()
            if not isinstance(uploaded, dict):
                raise RuntimeError("Drive API upload response was not a dictionary.")

            uploaded_id = uploaded.get("id")
            if not uploaded_id:
                raise RuntimeError(
                    f"Drive upload for '{doc_title}' succeeded without returning a file ID."
                )

            link = uploaded.get("webViewLink", "")
            if not link:
                # Build a stable editor URL when Drive omits webViewLink.
                link = f"https://docs.google.com/document/d/{uploaded_id}/edit"

            log_docs_sync(
                "INFO",
                "Document uploaded and converted to Google Doc.",
                doc_title=doc_title,
                source_path=docx_path,
                target_folder_id=folder_id,
                uploaded_doc_id=uploaded_id,
                web_link=link
            )
            return link

        # Upload all three documents into the Assignments folder
        brief_link  = upload_docx_as_google_doc(
            brief_path,
            f"[Brief] {brief.title} [{assignment_id}]",
            target_id
        )
        rubric_link = upload_docx_as_google_doc(
            rubric_path,
            f"[Rubric] {rubric.title} [{assignment_id}]",
            target_id
        )
        scheme_link = upload_docx_as_google_doc(
            scheme_path,
            f"[Scheme] {scheme.title} [{assignment_id}]",
            target_id
        )

        log_docs_sync(
            "INFO",
            "All three documents uploaded to /AkadVerse/2026/Assignments/ successfully.",
            target_folder_id=target_id,
            brief_link=brief_link,
            rubric_link=rubric_link,
            scheme_link=scheme_link
        )
        return True

    except ImportError:
        log_docs_sync(
            "WARNING",
            "google-api-python-client not installed. "
            "Run: pip install google-api-python-client google-auth"
        )
        return False
    except Exception as e:
        # Docs sync failure is non-fatal — log clearly and return False
        log_docs_sync(
            "WARNING",
            "Sync failed. Assignment still saved locally.",
            error_type=type(e).__name__,
            error=str(e)
        )
        return False


# =========================================================
# GENERATION PROMPT
# =========================================================

GENERATION_PROMPT = PromptTemplate(
    template="""You are an experienced Nigerian university lecturer and academic assessment designer.

Generate a complete, professional assignment package for the following course and topic.
All three documents must be consistent with each other — total marks, tasks, and criteria
must match across the brief, rubric, and marking scheme.

COURSE DETAILS:
  Course Title:       {course_title}
  Course Code:        {course_code}
  Topic:              {topic}
  Academic Level:     {academic_level}
  Total Marks:        {total_marks}
  Weighting:          {weighting}% of final grade
  Submission Deadline: {deadline}
  Assignment Type:    {assignment_type}

LEARNING OUTCOMES TO ASSESS:
{learning_outcomes}

ADDITIONAL INSTRUCTIONS FROM LECTURER:
{additional_instructions}

REQUIREMENTS:
- The assignment brief must be clear, unambiguous, and appropriate for the stated level.
- The rubric must have 4-6 criteria covering the key aspects of the assignment.
  Each criterion must have distinct descriptors for Distinction, Merit, Pass, and Fail.
  The marks across all criteria must sum to exactly {total_marks}.
- The marking scheme must provide thorough model answers for every task in the brief.
  Mark allocations in the scheme must also sum to exactly {total_marks}.
- Use language and academic conventions appropriate for Nigerian universities.
- The academic integrity note should reference the institution's standard policy.

Respond with a single JSON object matching the AssignmentPackage schema exactly.
""",
    input_variables=[
        "course_title", "course_code", "topic", "academic_level",
        "total_marks", "weighting", "deadline", "assignment_type",
        "learning_outcomes", "additional_instructions"
    ]
)


# =========================================================
# FASTAPI APPLICATION
# =========================================================

@asynccontextmanager
async def lifespan(_: "FastAPI") -> AsyncIterator[None]:
    """
    Runs on startup: initialises the database and ensures the output
    directory exists. Runs on shutdown: logs a clean stop message.
    """
    print("[Startup] AkadVerse Assignment Generator initialising...")

    # Ensure the local output directory exists
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"[Startup] Output directory: '{OUTPUT_DIR}'")

    # Initialise SQLite metadata table
    init_db()

    print("[Startup] Ready. Run with: uvicorn assignment_generator:app --host 127.0.0.1 --port 8008 --reload")
    yield
    print("[Shutdown] AkadVerse Assignment Generator stopped.")


app = FastAPI(
    title="AkadVerse — Assignment Generator API",
    description=(
        "Tier 5 faculty tool. Generates a complete assignment package "
        "(brief, marking rubric, marking scheme) from course context using Gemini."
    ),
    version="1.0",
    lifespan=lifespan
)


# =========================================================
# ENDPOINT 1: Generate a complete assignment package
# =========================================================

@app.post("/generate-assignment", response_model=GenerationSuccessResponse, tags=["Generation"])
async def generate_assignment(
    course_title: str = Form(..., description="Full course title, e.g. 'Data Structures and Algorithms'."),
    course_code: str = Form(..., description="Course code, e.g. 'CSC 301'."),
    topic: str = Form(..., description="Specific topic the assignment covers, e.g. 'Binary Search Trees'."),
    academic_level: str = Form(..., description="Academic level, e.g. '300 Level', 'Year 2', 'Postgraduate'."),
    total_marks: int = Form(..., ge=10, le=100, description="Total marks for the assignment (10-100)."),
    weighting: int = Form(..., ge=5, le=60, description="Percentage of final grade this assignment contributes (5-60%)."),
    deadline: str = Form(..., description="Submission deadline, e.g. 'Two weeks from date of issue'."),
    assignment_type: str = Form(
        default="Individual Assignment",
        description="Type of assignment, e.g. 'Individual Assignment', 'Group Project', 'Lab Report'."
    ),
    learning_outcomes: str = Form(
        ...,
        description="The learning outcomes this assignment assesses. List them clearly, one per line."
    ),
    additional_instructions: str = Form(
        default="None.",
        description="Any specific instructions for Gemini, e.g. 'Include a programming task in Python', 'Focus on theoretical analysis'."
    ),
    sync_to_docs: bool = Form(
        default=False,
        description="If True, attempts to sync the generated documents to Google Docs using Tier 4 credentials."
    ),
    google_api_key: str = Form(..., description="Your Google Gemini API key.")
):
    """
    Generates a complete assignment package in one Gemini API call.

    The package contains three documents:
      1. Assignment Brief  — student-facing task description
      2. Marking Rubric    — grade band descriptors per criterion
      3. Marking Scheme    — model answers and mark allocations

    All three are saved as .docx files locally (simulating GCS) and
    their metadata is logged to SQLite (simulating PostgreSQL).
    Optionally syncs to Google Docs if Tier 4 credentials are present.
    """

    # -- Input validation --
    if len(learning_outcomes.strip()) < 20:
        raise HTTPException(
            status_code=400,
            detail="Learning outcomes are too short. Please provide meaningful outcomes."
        )
    if len(learning_outcomes) > MAX_LEARNING_OUTCOMES_LENGTH:
        raise HTTPException(
            status_code=400,
            detail=f"Learning outcomes text exceeds {MAX_LEARNING_OUTCOMES_LENGTH} characters. Please be concise."
        )

    # -- Step 1: Select the generative model dynamically --
    selected_model = get_valid_model_name(google_api_key)

    llm = ChatGoogleGenerativeAI(
        model=selected_model,
        api_key=google_api_key,
        temperature=0.4   # Moderate temperature: creative enough for good writing,
                          # grounded enough for consistent mark totals
    )

    # Bind the LLM to our AssignmentPackage schema for structured output
    structured_llm = llm.with_structured_output(AssignmentPackage, method="json_mode")

    # -- Step 2: Format the prompt --
    prompt_text = GENERATION_PROMPT.format(
        course_title=course_title,
        course_code=course_code,
        topic=topic,
        academic_level=academic_level,
        total_marks=total_marks,
        weighting=weighting,
        deadline=deadline,
        assignment_type=assignment_type,
        learning_outcomes=learning_outcomes.strip(),
        additional_instructions=additional_instructions.strip() or "None."
    )

    # -- Step 3: Generate the assignment package --
    print(f"[Generate] Generating assignment: '{topic}' for {course_code} ({academic_level})...")
    print(f"[Generate] Model: {selected_model} | Marks: {total_marks} | Type: {assignment_type}")

    try:
        package: AssignmentPackage = structured_llm.invoke(prompt_text)  # type: ignore
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Gemini generation failed: {e}"
        )

    print(f"[Generate] Package generated successfully. Saving documents...")

    # -- Step 4: Build file paths using a unique assignment ID --
    assignment_id = uuid4().hex[:12]
    # Sanitise the topic for use in filenames (remove special chars)
    safe_topic = "".join(c if c.isalnum() or c in " _-" else "" for c in topic).strip().replace(" ", "_")
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    file_prefix = f"{assignment_id}_{safe_topic}_{timestamp}"

    brief_path  = os.path.join(OUTPUT_DIR, f"{file_prefix}_brief.docx")
    rubric_path = os.path.join(OUTPUT_DIR, f"{file_prefix}_rubric.docx")
    scheme_path = os.path.join(OUTPUT_DIR, f"{file_prefix}_scheme.docx")

    # -- Step 5: Save the three .docx files --
    try:
        save_brief_as_docx(package.brief, brief_path)
        save_rubric_as_docx(package.rubric, rubric_path)
        save_scheme_as_docx(package.scheme, scheme_path)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to save generated documents: {e}"
        )

    # -- Step 6: Optionally sync to Google Docs --
    docs_synced = False
    if sync_to_docs:
        print("[Generate] Attempting Google Docs sync...")
        docs_synced = sync_to_google_docs(
            package.brief, package.rubric, package.scheme,
            assignment_id, google_api_key,
            # Pass the actual saved file paths so the sync function
            # uploads the real .docx files rather than recreating them
            brief_path=brief_path,
            rubric_path=rubric_path,
            scheme_path=scheme_path
        )

    # -- Step 7: Log metadata to SQLite --
    try:
        log_assignment_to_db(
            assignment_id=assignment_id,
            course_title=course_title,
            topic=topic,
            academic_level=academic_level,
            total_marks=total_marks,
            brief_path=brief_path,
            rubric_path=rubric_path,
            scheme_path=scheme_path,
            docs_synced=docs_synced
        )
    except Exception as e:
        # DB logging failure is non-fatal — files are already saved
        print(f"[DB WARNING] Metadata logging failed: {e}. Files are saved but not recorded in DB.")

    # -- Step 8: Publish Kafka mock event --
    print(f"[KAFKA MOCK] Published event 'assignment.created' — ID: {assignment_id}, Course: {course_code}, Topic: {topic}")

    return GenerationSuccessResponse(
        status="success",
        assignment_id=assignment_id,
        message=(
            f"Assignment package for '{topic}' generated successfully. "
            f"Three documents saved to '{OUTPUT_DIR}'."
        ),
        course_title=course_title,
        topic=topic,
        files_saved=[brief_path, rubric_path, scheme_path],
        docs_synced=docs_synced,
        total_marks=package.brief.total_marks
    )


# =========================================================
# ENDPOINT 2: List previously generated assignments
# =========================================================

@app.get("/assignments", tags=["History"])
async def list_assignments(
    course_title: Optional[str] = None,
    limit: int = 10,
    offset: int = 0
):
    """
    Returns a paginated list of previously generated assignments.
    Optionally filter by course title (partial match).

    Useful for a lecturer who wants to review what they have generated
    without opening the file system.
    """
    try:
        with get_db_connection() as conn:
            if course_title:
                rows = conn.execute(
                    """
                    SELECT id, course_title, topic, academic_level, total_marks,
                           docs_synced, created_at
                    FROM assignments
                    WHERE course_title LIKE ?
                    ORDER BY created_at DESC
                    LIMIT ? OFFSET ?
                    """,
                    (f"%{course_title}%", limit, offset)
                ).fetchall()
            else:
                rows = conn.execute(
                    """
                    SELECT id, course_title, topic, academic_level, total_marks,
                           docs_synced, created_at
                    FROM assignments
                    ORDER BY created_at DESC
                    LIMIT ? OFFSET ?
                    """,
                    (limit, offset)
                ).fetchall()

        if not rows:
            return {"assignments": [], "message": "No assignments found."}

        return {
            "assignments": [
                {
                    "id": row["id"],
                    "course_title": row["course_title"],
                    "topic": row["topic"],
                    "academic_level": row["academic_level"],
                    "total_marks": row["total_marks"],
                    "docs_synced": bool(row["docs_synced"]),
                    "created_at": row["created_at"]
                }
                for row in rows
            ],
            "total_returned": len(rows)
        }

    except sqlite3.Error as e:
        raise HTTPException(status_code=500, detail=f"Database error: {e}")


# =========================================================
# ENDPOINT 3: Health check
# =========================================================

@app.get("/health", tags=["System"])
async def health_check():
    """Reports service status and counts of stored assignments."""
    try:
        with get_db_connection() as conn:
            count = conn.execute("SELECT COUNT(*) FROM assignments").fetchone()[0]
    except Exception:
        count = 0

    return {
        "status": "ok",
        "version": "1.0",
        "output_directory": OUTPUT_DIR,
        "output_dir_exists": os.path.exists(OUTPUT_DIR),
        "assignments_generated": count,
        "docs_sync_available": os.path.exists("token.json"),
        "endpoints": {
            "POST /generate-assignment": "Generate a full assignment package (brief + rubric + scheme).",
            "GET  /assignments": "List previously generated assignments with optional filtering.",
            "GET  /health": "This endpoint."
        }
    }


# =========================================================
# Run: uvicorn assignment_generator:app --host 127.0.0.1 --port 8008 --reload
# =========================================================